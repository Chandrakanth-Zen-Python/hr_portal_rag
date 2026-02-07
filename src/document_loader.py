from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List

from PyPDF2 import PdfReader


@dataclass
class Document:
    text: str
    source: str


def _load_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)


def _load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _load_docx(path: Path) -> str:
    try:
        import docx
    except Exception as exc:  # pragma: no cover - optional dependency
        raise RuntimeError(
            "python-docx is required to load .docx files. Install it with pip."
        ) from exc

    doc = docx.Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def load_documents(paths: Iterable[str]) -> List[Document]:
    docs: List[Document] = []
    for raw in paths:
        path = Path(raw)
        if not path.exists():
            continue
        ext = path.suffix.lower()
        if ext == ".pdf":
            text = _load_pdf(path)
        elif ext == ".docx":
            text = _load_docx(path)
        else:
            text = _load_text(path)

        if text:
            text = text.replace("\x00", "")
        if text.strip():
            docs.append(Document(text=text, source=str(path)))
    return docs


def load_from_directory(directory: str) -> List[Document]:
    root = Path(directory)
    files = [str(p) for p in root.rglob("*") if p.is_file()]
    return load_documents(files)
